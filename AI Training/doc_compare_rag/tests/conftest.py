"""
conftest.py – Session-scoped fixtures shared across all test modules.

A real .docx is built with python-docx; a real text PDF is built with fpdf2.
Both are created once per test session and reused.

Pytesseract and pdf2image are added to sys.modules as MagicMocks at import time
so tests can patch their attributes even when the binaries are not installed.
"""

import sys
from unittest.mock import MagicMock

import pytest

# ---------------------------------------------------------------------------
# Stub out optional binary-dependent packages so patch() can intercept them.
# ---------------------------------------------------------------------------

for _pkg in ("pytesseract", "pdf2image", "sentence_transformers"):
    if _pkg not in sys.modules:
        sys.modules[_pkg] = MagicMock()


# ---------------------------------------------------------------------------
# .docx fixture
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def minimal_docx_path(tmp_path_factory):
    """Minimal .docx with two headings and body paragraphs."""
    from docx import Document as DocxDocument  # type: ignore

    tmp = tmp_path_factory.mktemp("fixtures")
    path = tmp / "sample.docx"

    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction text. It covers the basics.")
    doc.add_paragraph("A second paragraph under the introduction heading.")
    doc.add_heading("Background", level=1)
    doc.add_paragraph("Background information about the subject matter.")
    doc.save(str(path))
    return str(path)


# ---------------------------------------------------------------------------
# Text PDF fixture
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def minimal_pdf_path(tmp_path_factory):
    """Minimal single-page text PDF built with fpdf2."""
    from fpdf import FPDF  # type: ignore

    tmp = tmp_path_factory.mktemp("fixtures")
    path = tmp / "sample.pdf"

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "INTRODUCTION", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(
        0,
        10,
        "This is a sample PDF document created for testing purposes.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(
        0,
        10,
        "It contains enough text to exceed the fifty-character threshold.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.output(str(path))
    return str(path)
